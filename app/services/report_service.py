from datetime import datetime, time, timedelta
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import HTTPException
from app.models.behavior_record import BehaviorRecord
from app.models.student import Student
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def get_behavior_class_report_data(db, current_user, data):
    if not current_user.school_id:
        raise HTTPException(status_code=403, detail="User is not linked to a school")

    class_letter = data.class_letter.strip().upper()
    date_from = datetime.combine(data.date_from, time.min)
    date_to = datetime.combine(data.date_to, time.min) + timedelta(days=1)

    records = (
        db.query(BehaviorRecord, Student)
        .join(Student, Student.id == BehaviorRecord.student_id)
        .filter(
            BehaviorRecord.school_id == current_user.school_id,
            Student.school_id == current_user.school_id,
            Student.grade == data.grade,
            Student.class_letter == class_letter,
            BehaviorRecord.created_at >= date_from,
            BehaviorRecord.created_at < date_to,
        )
        .order_by(
            BehaviorRecord.created_at.desc(), Student.last_name, Student.first_name
        )
        .all()
    )

    items = []
    for record, student in records:
        full_name = " ".join(
            part
            for part in (student.last_name, student.first_name, student.middle_name)
            if part
        )
        violation = ", ".join(record.reasons) if record.reasons else ""

        items.append(
            {
                "full_name": full_name,
                "class_name": f"{student.grade}{student.class_letter}",
                "subject": record.subject,
                "date": record.created_at.date(),
                "violation": violation,
            }
        )

    return {
        "school_id": current_user.school_id,
        "grade": data.grade,
        "class_letter": class_letter,
        "date_from": data.date_from,
        "date_to": data.date_to,
        "total": len(items),
        "items": items,
    }


def build_behavior_excel(report_data):
    output = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Поведение"

    title = (
        f"Отчет по поведению за период "
        f"{report_data['date_from'].strftime('%d.%m.%Y')} - "
        f"{report_data['date_to'].strftime('%d.%m.%Y')}"
    )
    class_info = f"Класс: {report_data['grade']}{report_data['class_letter']}"

    sheet.merge_cells("A1:E1")
    sheet["A1"] = title
    sheet["A1"].font = Font(bold=True, size=14)
    sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")

    sheet.merge_cells("A2:E2")
    sheet["A2"] = class_info
    sheet["A2"].font = Font(bold=True, size=11)
    sheet["A2"].alignment = Alignment(horizontal="left", vertical="center")

    headers = ["ФИО", "Класс", "Предмет", "Дата", "Нарушение"]
    header_row = 4

    for col_index, header in enumerate(headers, start=1):
        cell = sheet.cell(row=header_row, column=col_index, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    thin_side = Side(style="thin", color="BFBFBF")
    border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

    for col in range(1, 6):
        sheet.cell(row=header_row, column=col).border = border

    data_start_row = header_row + 1
    current_row = data_start_row

    for item in report_data["items"]:
        values = [
            item["full_name"],
            item["class_name"],
            item["subject"],
            item["date"].strftime("%d.%m.%Y"),
            item["violation"],
        ]

        for col_index, value in enumerate(values, start=1):
            cell = sheet.cell(row=current_row, column=col_index, value=value)
            cell.border = border
            cell.alignment = Alignment(
                horizontal="left" if col_index in (1, 3, 5) else "center",
                vertical="top",
                wrap_text=True,
            )

        current_row += 1

    sheet.freeze_panes = "A5"
    sheet.sheet_view.showGridLines = False

    column_widths = {
        "A": 34,
        "B": 12,
        "C": 22,
        "D": 14,
        "E": 60,
    }
    for column, width in column_widths.items():
        sheet.column_dimensions[column].width = width

    sheet.row_dimensions[1].height = 24
    sheet.row_dimensions[2].height = 20
    sheet.row_dimensions[4].height = 22

    workbook.save(output)
    output.seek(0)
    return output


def build_behavior_docx(report_data):
    output = BytesIO()
    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(
        "Отчет по поведению за период "
        f"{report_data['date_from'].strftime('%d.%m.%Y')} - "
        f"{report_data['date_to'].strftime('%d.%m.%Y')}"
    )
    title_run.bold = True
    title_run.font.size = Pt(14)

    class_info = document.add_paragraph(
        f"Класс: {report_data['grade']}{report_data['class_letter']}"
    )
    class_info.runs[0].bold = True

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["ФИО", "Класс", "Предмет", "Дата", "Нарушение"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    for item in report_data["items"]:
        row = table.add_row().cells
        row[0].text = item["full_name"]
        row[1].text = item["class_name"]
        row[2].text = item["subject"]
        row[3].text = item["date"].strftime("%d.%m.%Y")
        row[4].text = item["violation"]

    document.save(output)
    output.seek(0)
    return output


def build_behavior_pdf(report_data):
    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )

    styles = getSampleStyleSheet()
    story = []

    title = (
        "Отчет по поведению за период "
        f"{report_data['date_from'].strftime('%d.%m.%Y')} - "
        f"{report_data['date_to'].strftime('%d.%m.%Y')}"
    )
    story.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
    story.append(
        Paragraph(
            f"<b>Класс:</b> {report_data['grade']}{report_data['class_letter']}",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 8))

    table_data = [["ФИО", "Класс", "Предмет", "Дата", "Нарушение"]]
    for item in report_data["items"]:
        table_data.append(
            [
                item["full_name"],
                item["class_name"],
                item["subject"],
                item["date"].strftime("%d.%m.%Y"),
                item["violation"],
            ]
        )

    table = Table(
        table_data,
        repeatRows=1,
        colWidths=[70 * mm, 20 * mm, 38 * mm, 25 * mm, 120 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFBFBF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 1),  (1, -1), "CENTER"),
                ("ALIGN", (3, 1),  (3, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)

    doc.build(story)
    output.seek(0)
    return output
